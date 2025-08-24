import traceback
import shutil

import pandas as pd
import pythoncom
from docx import Document
from pathlib import Path
# from doc2docx import convert
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.shared import Pt


def change_unloading_file(incoming_data: dict, current_progress, now_doc: int, all_doc: int,
                          line_doing, line_progress, progress_value, event, window_check, info_value) -> dict:
    logging = incoming_data['logging']
    errors = []
    try:
        percent = 100/all_doc
        for file in Path(incoming_data['start_path']).rglob('*.*'):
            if window_check.stop_threading:
                return {'status': 'cancel', 'trace': '', 'text': ''}
            if '~' in file.name:
                continue
            if file.suffix not in ['.doc', '.docx']:
                shutil.copy(str(file), incoming_data['finish_path'])
                logging.info(f"Скопировали файл {file.name}")
                continue
            now_doc += 1
            line_doing.emit(f'Преобразуем файл {str(file.name)} ({now_doc} из {all_doc})')
            start_file = file
            finish_file = Path(incoming_data['finish_path'], file.stem + '.docx')
            if file.suffix == '.doc':
                errors.append(f"Файл {file.name} старого формата (.doc)")
                logging.warning(f"Файл {file.name} старого формата (.doc)")
                continue
                # try:
                #     pythoncom.CoInitializeEx(0)
                #     convert(str(start_file), str(finish_file))
                #     start_file = finish_file
                #     logging.info(f"Преобразовали файл {file.name}")
                # except BaseException as exception:
                #     errors.append(f"Ошибка при конвертации файла {file.name}")
                #     logging.warning(f"Ошибка при конвертации файла {file.name} - {exception}")
                #     continue
            document = Document(str(start_file))
            table = document.tables[0]
            number_columns = len(table.columns)
            if number_columns < 8:
                errors.append(f"Кол-во столбцов в таблице документа {file.name} меньше 8 (не та форма выгрузки)")
                logging.warning(f"Кол-во столбцов в таблице документа {file.name} меньше 8 (не та форма выгрузки)")
                continue
            df = pd.DataFrame()
            for index in range(len(table.columns)):
                df[index] = list(map(lambda val: val.text, table.column_cells(index)))
            change_index = df.loc[(df[2].str.len() == 0) & (df[3].str.len() == 0) & (df[5].str.len() == 0)
                                  & (df[6].str.len() == 0) & (df[7].str.len() == 0)].index.to_list()
            for row in change_index:
                name = table.cell(row, 1).text
                sn = table.cell(row, 4).text
                table.cell(row, 1).merge(table.cell(row, number_columns - 1))
                text = f"{name}, с/н {sn}" if len(sn) else f"{name}"
                table.cell(row, 1).text = text
                table.cell(row, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in table.cell(row, 1).paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        run.font.name = 'Times New Roman'
            document.save(str(finish_file))
            logging.info(f"Файл {file.name} успешно преобразован")
            current_progress += percent
            line_progress.emit(f'Выполнено {int(current_progress)} %')
            progress_value.emit(int(current_progress))
        return {'status': 'warning' if errors else 'success', 'text': errors, 'trace': ''}
    except BaseException as ex:
        return {'status': 'error', 'text': ex, 'trace': traceback.format_exc()}